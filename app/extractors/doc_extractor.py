import os
import logging
import subprocess
import tempfile
import docx

logger = logging.getLogger(__name__)

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from DOC/DOCX files
    
    Args:
        file_path: Path to the DOC/DOCX file
        
    Returns:
        Extracted text from the document
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.doc':
        # For older .doc files, convert to .docx first
        return extract_text_from_doc_legacy(file_path)
    else:
        logger.warning(f"Unsupported document format: {file_extension}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files using python-docx
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text
    """
    try:
        logger.info(f"Extracting text from DOCX: {file_path}")
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return '\n'.join(full_text)
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_doc_legacy(file_path: str) -> str:
    """
    Extract text from legacy DOC files by converting to DOCX using LibreOffice
    
    Args:
        file_path: Path to the DOC file
        
    Returns:
        Extracted text
    """
    try:
        logger.info(f"Converting legacy DOC to DOCX: {file_path}")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert .doc to .docx using LibreOffice
            output_file = os.path.join(temp_dir, "converted.docx")
            
            # Run LibreOffice to convert the file
            cmd = [
                'libreoffice', 
                '--headless', 
                '--convert-to', 
                'docx', 
                '--outdir', 
                temp_dir, 
                file_path
            ]
            
            process = subprocess.run(
                cmd, 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                timeout=30
            )
            
            if process.returncode != 0:
                logger.error(f"Error converting DOC to DOCX: {process.stderr.decode()}")
                return ""
            
            # Find the converted file
            for filename in os.listdir(temp_dir):
                if filename.endswith('.docx'):
                    converted_path = os.path.join(temp_dir, filename)
                    # Now extract text from the converted DOCX
                    return extract_text_from_docx(converted_path)
            
            logger.error("Converted DOCX file not found")
            return ""
    
    except Exception as e:
        logger.error(f"Error processing legacy DOC file: {str(e)}")
        return ""